"""章节 section_path：提取、切分继承、管线写入与 DOCX Heading。"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from app.rag.document_pipeline.enrichers import make_chunk_meta
from app.rag.document_pipeline.pipeline import DocumentPipeline
from app.rag.document_pipeline.section_utils import (
    docx_style_to_markdown_heading,
    normalize_heading_title,
    parse_heading_line,
    section_path_before_offset,
)
from app.rag.document_pipeline.splitters import ChunkingConfig, StructureSplitter, WindowSplitter
from app.rag.models import DocumentSource
from app.llm.graphs.chatbot_rag_citations import chunks_to_rag_citations
from app.rag.models import RetrievedChunk


class TestSectionUtils(unittest.TestCase):
    def test_markdown_heading(self) -> None:
        hm = parse_heading_line("## 二、现行技术标准与规范体系")
        self.assertIsNotNone(hm)
        assert hm is not None
        self.assertEqual(hm.level, 2)
        self.assertEqual(hm.section_path, "二、现行技术标准与规范体系")

    def test_numbered_heading(self) -> None:
        hm = parse_heading_line("1.2.3 监测方法")
        self.assertIsNotNone(hm)
        assert hm is not None
        self.assertEqual(hm.level, 3)
        self.assertEqual(hm.section_path, "1.2.3 监测方法")

    def test_cn_chapter(self) -> None:
        hm = parse_heading_line("第一章 总则")
        self.assertIsNotNone(hm)
        assert hm is not None
        self.assertEqual(hm.level, 1)
        self.assertIn("第一章", hm.section_path)
        self.assertIn("总则", hm.section_path)

    def test_cn_enum(self) -> None:
        hm = parse_heading_line("一、背景与危害概述")
        self.assertIsNotNone(hm)
        assert hm is not None
        self.assertEqual(hm.section_path, "一、背景与危害概述")

    def test_reject_long_or_list_like(self) -> None:
        self.assertIsNone(parse_heading_line("1. "))
        self.assertIsNone(parse_heading_line("x" * 100))

    def test_normalize_strips_hashes_and_truncates(self) -> None:
        self.assertEqual(normalize_heading_title("###  标题  "), "标题")
        long_t = "字" * 200
        out = normalize_heading_title(long_t)
        self.assertLessEqual(len(out), 120)
        self.assertTrue(out.endswith("…"))

    def test_docx_style_mapping(self) -> None:
        self.assertEqual(docx_style_to_markdown_heading("Heading 1", "引言"), "# 引言")
        self.assertEqual(docx_style_to_markdown_heading("标题 2", "方法"), "## 方法")
        self.assertIsNone(docx_style_to_markdown_heading("Normal", "正文"))

    def test_section_path_before_offset(self) -> None:
        text = "## A\nbody1\n## B\nbody2"
        path, level = section_path_before_offset(text, text.index("body2"))
        self.assertEqual(path, "B")
        self.assertEqual(level, 2)


class TestStructureSplitterSections(unittest.TestCase):
    def test_split_sections_assigns_path(self) -> None:
        text = (
            "## 一、背景\n"
            "地面沉降危害巨大。\n"
            "## 二、标准\n"
            "已建立多层次标准体系。"
        )
        blocks = StructureSplitter().split_sections(text)
        self.assertGreaterEqual(len(blocks), 2)
        paths = [b.section_path for b in blocks]
        self.assertIn("一、背景", paths)
        self.assertIn("二、标准", paths)
        # 各节正文非空
        self.assertTrue(all(b.text.strip() for b in blocks))

    def test_preamble_without_heading_has_null_path(self) -> None:
        text = "序言段落无标题。\n## 第一章\n正文。"
        blocks = StructureSplitter().split_sections(text)
        self.assertGreaterEqual(len(blocks), 2)
        self.assertIsNone(blocks[0].section_path)
        self.assertEqual(blocks[1].section_path, "第一章")


class TestDocumentPipelineSectionPath(unittest.TestCase):
    def _pipeline(self, strategy: str = "structure") -> DocumentPipeline:
        cfg = ChunkingConfig(chunk_size=80, chunk_overlap=10, min_chunk_size=5)
        return DocumentPipeline(cfg=cfg, strategy=strategy, cleaning_profile="normal")

    def test_structure_pipeline_writes_section_path(self) -> None:
        content = (
            "## 一、背景与危害概述\n"
            "地面沉降是我国面临的一项严峻地质灾害，目前已有多个城市遭受影响。\n\n"
            "## 二、现行技术标准与规范体系\n"
            "为规范地面沉降的调查、监测与防治工作，我国已建立起多层次的技术标准体系。"
        )
        src = DocumentSource(
            dataset_id="t",
            doc_name="沉降综述",
            namespace="kb",
            content=content,
            source_type="markdown",
        )
        with patch("app.rag.document_pipeline.pipeline.get_app_config") as mock_cfg:
            mock_cfg.return_value.rag.ingestion.chunk_size = 80
            mock_cfg.return_value.rag.ingestion.chunk_overlap = 10
            mock_cfg.return_value.rag.ingestion.min_chunk_size = 5
            mock_cfg.return_value.rag.ingestion.default_chunk_strategy = "structure"
            mock_cfg.return_value.rag.ingestion.cleaning_profile = "normal"
            mock_cfg.return_value.rag.ingestion.clean_remove_header_footer = True
            mock_cfg.return_value.rag.ingestion.clean_merge_duplicate_paragraphs = True
            mock_cfg.return_value.rag.ingestion.clean_fix_encoding_noise = True
            mock_cfg.return_value.rag.ingestion.clean_strip_html = True
            mock_cfg.return_value.rag.ingestion.clean_min_repeated_line_pages = 2
            pipe = self._pipeline("structure")
            staged = pipe.process_document_staged(src)

        chunks = staged["chunks"]
        self.assertGreaterEqual(len(chunks), 2)
        paths = {c.metadata.get("section_path") for c in chunks}
        self.assertTrue(any(p and "背景" in str(p) for p in paths))
        self.assertTrue(any(p and "标准" in str(p) for p in paths))
        self.assertGreaterEqual(int(staged["stats"].get("chunks_with_section_path") or 0), 1)
        # 同章节子块继承
        for c in chunks:
            if c.metadata.get("section_path"):
                self.assertIn("section_level", c.metadata)

    def test_no_heading_section_path_absent(self) -> None:
        src = DocumentSource(
            dataset_id="t",
            doc_name="plain",
            namespace="kb",
            content="这是一段没有标题的普通正文，用于验证 section_path 可为空。",
            source_type="text",
        )
        with patch("app.rag.document_pipeline.pipeline.get_app_config") as mock_cfg:
            mock_cfg.return_value.rag.ingestion.chunk_size = 500
            mock_cfg.return_value.rag.ingestion.chunk_overlap = 80
            mock_cfg.return_value.rag.ingestion.min_chunk_size = 40
            mock_cfg.return_value.rag.ingestion.default_chunk_strategy = "structure"
            mock_cfg.return_value.rag.ingestion.cleaning_profile = "normal"
            mock_cfg.return_value.rag.ingestion.clean_remove_header_footer = True
            mock_cfg.return_value.rag.ingestion.clean_merge_duplicate_paragraphs = True
            mock_cfg.return_value.rag.ingestion.clean_fix_encoding_noise = True
            mock_cfg.return_value.rag.ingestion.clean_strip_html = True
            mock_cfg.return_value.rag.ingestion.clean_min_repeated_line_pages = 2
            pipe = self._pipeline("structure")
            chunks, _ = pipe.process_document(src)
        self.assertTrue(chunks)
        self.assertTrue(all(not c.metadata.get("section_path") for c in chunks))

    def test_window_strategy_annotates_nearest_heading(self) -> None:
        # 构造足够长、带标题的正文，迫使 window 切出多块
        body_a = "背景说明。" * 40
        body_b = "标准说明。" * 40
        content = f"## 甲章\n{body_a}\n## 乙章\n{body_b}"
        src = DocumentSource(
            dataset_id="t",
            doc_name="win",
            namespace="kb",
            content=content,
            source_type="markdown",
        )
        with patch("app.rag.document_pipeline.pipeline.get_app_config") as mock_cfg:
            mock_cfg.return_value.rag.ingestion.chunk_size = 120
            mock_cfg.return_value.rag.ingestion.chunk_overlap = 20
            mock_cfg.return_value.rag.ingestion.min_chunk_size = 10
            mock_cfg.return_value.rag.ingestion.default_chunk_strategy = "window"
            mock_cfg.return_value.rag.ingestion.cleaning_profile = "normal"
            mock_cfg.return_value.rag.ingestion.clean_remove_header_footer = True
            mock_cfg.return_value.rag.ingestion.clean_merge_duplicate_paragraphs = True
            mock_cfg.return_value.rag.ingestion.clean_fix_encoding_noise = True
            mock_cfg.return_value.rag.ingestion.clean_strip_html = True
            mock_cfg.return_value.rag.ingestion.clean_min_repeated_line_pages = 2
            pipe = DocumentPipeline(
                cfg=ChunkingConfig(chunk_size=120, chunk_overlap=20, min_chunk_size=10),
                strategy="window",
            )
            chunks, stats = pipe.process_document(src)
        self.assertGreaterEqual(len(chunks), 2)
        paths = {c.metadata.get("section_path") for c in chunks if c.metadata.get("section_path")}
        self.assertTrue(paths.intersection({"甲章", "乙章"}) or any("甲" in str(p) or "乙" in str(p) for p in paths))


class TestMakeChunkMetaSection(unittest.TestCase):
    def test_optional_section_fields(self) -> None:
        meta = make_chunk_meta("d", 0, "ns", None, section_path="  章A  ", section_level=2)
        self.assertEqual(meta["section_path"], "章A")
        self.assertEqual(meta["section_level"], 2)
        meta2 = make_chunk_meta("d", 1, "ns", None)
        self.assertNotIn("section_path", meta2)


class TestRagCitationsSectionPath(unittest.TestCase):
    def test_citations_include_section_path(self) -> None:
        chunks = [
            RetrievedChunk(
                text="片段正文",
                doc_name="沉降综述",
                namespace="kb",
                chunk_id="c1",
                section_path="二、现行技术标准与规范体系",
                metadata={"content_fetched_from_url": "https://example.com/a.pdf"},
            )
        ]
        cites = chunks_to_rag_citations(chunks)
        self.assertEqual(len(cites), 1)
        self.assertEqual(cites[0]["section_path"], "二、现行技术标准与规范体系")
        self.assertEqual(cites[0]["doc_name"], "沉降综述")
        self.assertEqual(cites[0]["original_content_url"], "https://example.com/a.pdf")


class TestDocxHeadingParse(unittest.TestCase):
    def test_docx_heading_becomes_markdown_and_section_path(self) -> None:
        try:
            import docx  # type: ignore[import-untyped]
        except Exception:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "sample.docx"
            document = docx.Document()
            document.add_heading("总则", level=1)
            document.add_paragraph("本章规定基本原则。")
            document.add_heading("监测要求", level=2)
            document.add_paragraph("应当按规范开展监测。")
            document.save(str(path))

            src = DocumentSource(
                dataset_id="t",
                doc_name="规范",
                namespace="kb",
                content=str(path),
                source_type="docx",
            )
            with patch("app.rag.document_pipeline.pipeline.get_app_config") as mock_cfg:
                mock_cfg.return_value.rag.ingestion.chunk_size = 200
                mock_cfg.return_value.rag.ingestion.chunk_overlap = 20
                mock_cfg.return_value.rag.ingestion.min_chunk_size = 10
                mock_cfg.return_value.rag.ingestion.default_chunk_strategy = "structure"
                mock_cfg.return_value.rag.ingestion.cleaning_profile = "normal"
                mock_cfg.return_value.rag.ingestion.clean_remove_header_footer = True
                mock_cfg.return_value.rag.ingestion.clean_merge_duplicate_paragraphs = True
                mock_cfg.return_value.rag.ingestion.clean_fix_encoding_noise = True
                mock_cfg.return_value.rag.ingestion.clean_strip_html = True
                mock_cfg.return_value.rag.ingestion.clean_min_repeated_line_pages = 2
                pipe = DocumentPipeline(
                    cfg=ChunkingConfig(chunk_size=200, chunk_overlap=20, min_chunk_size=10),
                    strategy="structure",
                )
                staged = pipe.process_document_staged(src)

            parsed = staged["parsed"]
            self.assertIn("# 总则", parsed)
            self.assertIn("## 监测要求", parsed)
            paths = {c.metadata.get("section_path") for c in staged["chunks"]}
            self.assertTrue(any(p and "总则" in str(p) for p in paths))
            self.assertTrue(any(p and "监测要求" in str(p) for p in paths))


class TestWindowSplitterWithSections(unittest.TestCase):
    def test_split_with_sections(self) -> None:
        text = "## S1\n" + ("内容甲。" * 30) + "\n## S2\n" + ("内容乙。" * 30)
        blocks = WindowSplitter(ChunkingConfig(chunk_size=80, chunk_overlap=10, min_chunk_size=5)).split_with_sections(
            text
        )
        self.assertGreaterEqual(len(blocks), 2)
        self.assertTrue(any(b.section_path for b in blocks))


if __name__ == "__main__":
    unittest.main()
