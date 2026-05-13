from __future__ import annotations

import os
import tempfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from app.inspection_v2.docx_rich_text import normalize_shading_fill, serialize_docx_for_inspection_v2
from app.inspection_v2.processing_units import split_docx_v2_by_processing_units
from app.models.inspection_extract import InspectionExtractRequest
from app.services.inspection_extract_service import InspectionExtractService


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shd)


def test_normalize_shading_fill_argb_and_hash() -> None:
    assert normalize_shading_fill("FFFF0000") == "FF0000"
    assert normalize_shading_fill("#C00000") == "C00000"


def test_serialize_docx_marks_shading_candidate() -> None:
    doc = Document()
    doc.add_paragraph("锅炉检修")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "右墙"
    table.cell(0, 1).text = "4.2"
    _set_cell_shading(table.cell(0, 1), "FF0000")

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        out = serialize_docx_for_inspection_v2(path, candidate_fills={"FF0000"})
        assert "锅炉检修" in out
        assert "[DOCX_V2_TABLE" in out
        assert "超标候选" in out
        assert "底纹=FF0000" in out
        assert "c1=" in out
    finally:
        Path(path).unlink(missing_ok=True)


def test_serialize_docx_marks_any_color_even_non_candidate() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "5.2"
    _set_cell_shading(table.cell(0, 0), "00FF00")

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        out = serialize_docx_for_inspection_v2(path, candidate_fills=set())
        assert "颜色标注" in out
        assert "底纹=00FF00" in out
        assert "超标候选" not in out
    finally:
        Path(path).unlink(missing_ok=True)


def test_serialize_docx_marks_run_font_color() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("4.73")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        out = serialize_docx_for_inspection_v2(path, candidate_fills=set())
        assert "颜色标注" in out
        assert "字体=FF0000" in out
    finally:
        Path(path).unlink(missing_ok=True)


def test_serialize_docx_collapses_uniform_title_row() -> None:
    """跨列表题被复制到每列时，首行折叠为一条，避免 r0 每列重复长标题。"""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    title = "高再入口蠕胀数据测量（51×4mm）"
    for j in range(3):
        table.cell(0, j).text = title
    table.cell(1, 0).text = "x"
    table.cell(1, 1).text = "y"
    table.cell(1, 2).text = "z"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        out = serialize_docx_for_inspection_v2(path, candidate_fills=set())
        assert "重复表题×3" in out
        assert out.count(title) == 1
    finally:
        Path(path).unlink(missing_ok=True)


def test_serialize_docx_corner_root_row_annotation() -> None:
    """斜线表角「根数+排数」同格时补充横纵语义，便于与管号列/行号列对齐。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "根数 排数"
    table.cell(0, 1).text = "1"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        out = serialize_docx_for_inspection_v2(path, candidate_fills=set())
        assert "[表角:" in out
        assert "横向表头=根数" in out
        assert "纵向表头=排数" in out
    finally:
        Path(path).unlink(missing_ok=True)


def test_serialize_docx_omits_default_black_font_marker() -> None:
    """显式默认黑 000000 不输出字体标注（避免与海量默认单元格重复 token）。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("3.4")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        out = serialize_docx_for_inspection_v2(path, candidate_fills=set())
        assert "字体=000000" not in out
        assert "颜色标注" not in out
    finally:
        Path(path).unlink(missing_ok=True)


def test_parse_document_v2_returns_docx_v2_route() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "x"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        svc = InspectionExtractService()  # type: ignore[call-arg]
        parsed, route = svc._parse_document_v2(  # noqa: SLF001
            InspectionExtractRequest(
                user_id="u1",
                session_id="s1",
                content=str(path),
                source_type="docx",
            )
        )
        assert route == "docx_v2"
        assert "[DOCX_V2_TABLE" in parsed
    finally:
        Path(path).unlink(missing_ok=True)


def test_docx_v2_single_chunk_per_logical_table() -> None:
    """整表必须在同一 chunk；不因 max_chunk_chars 较小而拆成多块。"""
    lines = [
        "1、测试小节标题",
        "[DOCX_V2_TABLE idx=1 rows=99 cols=2]",
        "r0: c0='列A' | c1='列B'",
        "r1: c0='说明' | c1='数值'",
    ]
    for i in range(2, 40):
        lines.append(f"r{i}: c0='位置{i}' | c1='{i}.5'")
    text = "\n".join(lines)
    chunks = split_docx_v2_by_processing_units(text, max_chunk_chars=320)
    table_chunks = [c for c in chunks if "[DOCX_V2_TABLE" in c]
    assert len(table_chunks) == 1
    assert table_chunks[0].count("r39:") == 1


def test_parse_document_v2_pdf_delegates_to_v1() -> None:
    svc = InspectionExtractService()  # type: ignore[call-arg]
    with patch.object(svc, "_parse_document", return_value=("pdf-body", "pdf_text")) as m:
        out, route = svc._parse_document_v2(  # noqa: SLF001
            InspectionExtractRequest(
                user_id="u1",
                session_id="s1",
                content="/tmp/x.pdf",
                source_type="pdf",
            )
        )
    assert (out, route) == ("pdf-body", "pdf_text")
    m.assert_called_once()
