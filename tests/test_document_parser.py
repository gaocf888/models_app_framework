from __future__ import annotations

from pathlib import Path

import pytest

from app.rag.document_pipeline.parsers import DocumentParser


def test_parse_docx_includes_table_in_document_order(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "sample.docx"
    d = Document()
    d.add_paragraph("第一段说明")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "类别"
    table.cell(0, 1).text = "表名"
    table.cell(1, 0).text = "设备"
    table.cell(1, 1).text = "account_boiler"
    d.add_paragraph("段落后文字")
    d.save(str(path))

    out = DocumentParser().parse(str(path), "docx")
    assert "第一段说明" in out
    assert "[DOCX_TABLE" in out
    assert "account_boiler" in out
    assert "段落后文字" in out
    # 表格应在两段文字之间（顺序）
    assert out.index("第一段说明") < out.index("[DOCX_TABLE")
    assert out.index("[DOCX_TABLE") < out.index("段落后文字")


def test_parse_xlsx_sheets_and_rows(tmp_path: Path) -> None:
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    path = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Meta"
    ws.append(["数据库表名", "中文描述"])
    ws.append(["account_boiler", "锅炉信息"])
    wb.save(str(path))

    out = DocumentParser().parse(str(path), "xlsx")
    assert "[XLSX_SHEET name=Meta]" in out
    assert "account_boiler" in out
    assert "锅炉信息" in out
